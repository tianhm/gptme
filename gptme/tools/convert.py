"""File format conversion tool for gptme.

Provides offline, multi-format file conversion with auto-detection and
graceful degradation when specific converters are unavailable.

Supported conversions:
  - PDF → PNG/JPEG (Poppler/pdftoppm, ImageMagick fallback)
  - Image → Image: PNG ↔ JPEG ↔ WebP (FFmpeg primary, ImageMagick fallback)
  - DOCX → text/markdown (python-docx, LibreOffice headless fallback)
  - MP4/video → JPEG thumbnail (FFmpeg)
  - PDF → text (pypdf, Poppler/pdftotext fallback)
"""

from __future__ import annotations

import glob as _glob
import importlib.util
import logging
import mimetypes
import shutil
import subprocess
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import ClassVar

from ..message import Message
from .base import Parameter, ToolSpec

logger = logging.getLogger(__name__)


def _arg(path: Path) -> str:
    """Render a path as a subprocess argument that cannot be read as an option.

    A file named e.g. ``-o`` would otherwise be parsed as a flag by tools that
    do not honour a ``--`` separator (pdftoppm, pdftotext). Prefixing ``./``
    keeps the path relative and identical, but no longer option-shaped.
    """
    text = str(path)
    return f"./{text}" if text.startswith("-") else text


def _im_src(path: Path) -> str:
    """Render a path safe for ImageMagick ``src[N]`` frame-selector syntax.

    ImageMagick interprets ``[`` and ``]`` in an input path as frame selectors,
    so a file named ``foo[0].pdf`` would silently match a different file. Escape
    any brackets in the path component before appending the ``[0]`` selector.
    """
    return _arg(path).replace("[", "\\[").replace("]", "\\]")


def _para_to_md(para: object) -> str:
    """Convert a python-docx paragraph to a Markdown line.

    Respects Heading 1/2/3 styles (→ ``#``/``##``/``###``) and list
    paragraphs (→ ``-``).  Inline bold and italic are rendered as
    ``**text**`` and ``*text*`` respectively via the paragraph's runs.
    Paragraphs with no text become empty lines (blank-line separators).
    """
    style_name: str = getattr(getattr(para, "style", None), "name", "") or ""
    runs = getattr(para, "runs", [])

    # Build inline text with bold/italic from runs; fall back to para.text.
    if runs:
        parts: list[str] = []
        for run in runs:
            text = run.text
            if not text:
                continue
            if getattr(run, "bold", False):
                text = f"**{text}**"
            elif getattr(run, "italic", False):
                text = f"*{text}*"
            parts.append(text)
        inline = "".join(parts)
    else:
        inline = getattr(para, "text", "")

    if not inline.strip():
        return ""

    if style_name.startswith("Heading 1"):
        return f"# {inline}"
    if style_name.startswith("Heading 2"):
        return f"## {inline}"
    if style_name.startswith("Heading 3"):
        return f"### {inline}"
    if "List" in style_name:
        return f"- {inline}"
    return inline


# ---------------------------------------------------------------------------
# Tool availability helpers
# ---------------------------------------------------------------------------


def _has_cmd(name: str) -> bool:
    return shutil.which(name) is not None


def _has_py(module: str) -> bool:
    return importlib.util.find_spec(module) is not None


@dataclass
class ToolAvailability:
    ffmpeg: bool = field(default_factory=lambda: _has_cmd("ffmpeg"))
    imagemagick: bool = field(default_factory=lambda: _has_cmd("convert"))
    pdftoppm: bool = field(default_factory=lambda: _has_cmd("pdftoppm"))
    pdftotext: bool = field(default_factory=lambda: _has_cmd("pdftotext"))
    libreoffice: bool = field(default_factory=lambda: _has_cmd("libreoffice"))
    tesseract: bool = field(default_factory=lambda: _has_cmd("tesseract"))
    python_magic: bool = field(default_factory=lambda: _has_py("magic"))
    python_docx: bool = field(default_factory=lambda: _has_py("docx"))
    pypdf: bool = field(default_factory=lambda: _has_py("pypdf"))

    def report(self) -> str:
        lines = []
        checks = [
            ("ffmpeg", "FFmpeg", self.ffmpeg),
            ("convert", "ImageMagick (convert)", self.imagemagick),
            ("pdftoppm", "Poppler (pdftoppm)", self.pdftoppm),
            ("pdftotext", "Poppler (pdftotext)", self.pdftotext),
            ("libreoffice", "LibreOffice headless", self.libreoffice),
            ("tesseract", "Tesseract OCR", self.tesseract),
            ("python-magic", "python-magic", self.python_magic),
            ("python-docx", "python-docx", self.python_docx),
            ("pypdf", "pypdf", self.pypdf),
        ]
        for _key, label, available in checks:
            mark = "✓" if available else "⚠"
            lines.append(f"  {mark} {label}")
        return "\n".join(lines)


_tool_avail: ToolAvailability | None = None


def get_availability() -> ToolAvailability:
    global _tool_avail
    if _tool_avail is None:
        _tool_avail = ToolAvailability()
    return _tool_avail


# ---------------------------------------------------------------------------
# Conversion result
# ---------------------------------------------------------------------------


@dataclass
class ConversionResult:
    success: bool
    output_path: Path | None
    converter_used: str
    warnings: list[str] = field(default_factory=list)
    error: str | None = None
    lossy: bool = False
    metadata: dict = field(default_factory=dict)

    def summary(self) -> str:
        if not self.success:
            return f"Conversion failed ({self.converter_used}): {self.error}"
        if self.metadata.get("dry_run"):
            return (
                f"Dry-run: would convert via {self.converter_used} → "
                f"{self.output_path} (no file written)"
            )
        parts = [f"Converted via {self.converter_used} → {self.output_path}"]
        if self.lossy:
            parts.append("(lossy)")
        if self.warnings:
            parts.append(f"Warnings: {'; '.join(self.warnings)}")
        return " | ".join(parts)


# ---------------------------------------------------------------------------
# Abstract converter base
# ---------------------------------------------------------------------------


class Converter(ABC):
    name: ClassVar[str] = "base"
    # Supported (source_mime_prefix, target_ext) pairs
    supported: ClassVar[list[tuple[str, str]]] = []

    @abstractmethod
    def is_available(self, avail: ToolAvailability) -> bool: ...

    @abstractmethod
    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult: ...

    def can_handle(self, src_mime: str, dest_ext: str) -> bool:
        dest_ext = dest_ext.lstrip(".").lower()
        for mime_prefix, ext in self.supported:
            if src_mime.startswith(mime_prefix) and ext == dest_ext:
                return True
        return False


# ---------------------------------------------------------------------------
# Concrete converters
# ---------------------------------------------------------------------------


class PDFToImageConverter(Converter):
    """PDF pages → raster images (PNG/JPEG) via pdftoppm or ImageMagick."""

    name = "pdf-to-image"
    supported = [
        ("application/pdf", "png"),
        ("application/pdf", "jpg"),
        ("application/pdf", "jpeg"),
    ]

    _DPI: ClassVar[dict[str, int]] = {"low": 72, "medium": 150, "high": 300}

    def is_available(self, avail: ToolAvailability) -> bool:
        return avail.pdftoppm or avail.imagemagick

    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult:
        avail = get_availability()
        dpi = self._DPI.get(quality, 150)
        dest_ext = dest.suffix.lstrip(".").lower()
        warnings: list[str] = []

        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used="",
                error=f"Destination is an existing directory: {dest}",
            )

        if avail.pdftoppm:
            # pdftoppm outputs <dest_stem>-<N>.ppm / .png; we want single page or first
            fmt_flag = "-png" if dest_ext == "png" else "-jpeg"
            with tempfile.TemporaryDirectory() as _tmpdir:
                tmp_prefix = Path(_tmpdir) / dest.stem
                cmd = [
                    "pdftoppm",
                    fmt_flag,
                    "-r",
                    str(dpi),
                    _arg(src),
                    _arg(tmp_prefix),
                ]
                result = subprocess.run(cmd, capture_output=True, check=False)
                if result.returncode == 0:
                    stem_esc = _glob.escape(dest.stem)
                    candidates = sorted(Path(_tmpdir).glob(f"{stem_esc}-*.{dest_ext}"))
                    if not candidates:
                        # Some versions use -1 suffix, or output format differs
                        candidates = sorted(
                            Path(_tmpdir).glob(
                                f"{stem_esc}*"
                                if dest_ext == "png"
                                else f"{stem_esc}*.jpg"
                            )
                        )
                    if candidates:
                        try:
                            shutil.move(str(candidates[0]), dest)
                        except OSError as e:
                            return ConversionResult(
                                success=False,
                                output_path=None,
                                converter_used="pdftoppm",
                                error=f"Could not write {dest}: {e}",
                                warnings=warnings,
                            )
                        # Guard against TOCTOU: if dest became a directory between
                        # the is_dir() check and shutil.move, the file lands inside
                        # it and dest.is_file() returns False.
                        if not dest.is_file():
                            return ConversionResult(
                                success=False,
                                output_path=None,
                                converter_used="pdftoppm",
                                error=f"Destination {dest} was replaced by a directory during conversion",
                                warnings=warnings,
                            )
                        if len(candidates) > 1:
                            warnings.append(
                                f"Multi-page PDF: only first page saved ({len(candidates)} pages total)"
                            )
                        return ConversionResult(
                            success=True,
                            output_path=dest,
                            converter_used="pdftoppm",
                            warnings=warnings,
                            lossy=dest_ext in ("jpg", "jpeg"),
                        )
                # tmpdir cleanup handles any partial page files automatically
            logger.warning(
                "pdftoppm failed: %s", result.stderr.decode(errors="replace")
            )
            warnings.append("pdftoppm failed; falling back to ImageMagick")

        if avail.imagemagick:
            density = str(dpi)
            cmd = [
                "convert",
                "-density",
                density,
                f"{_im_src(src)}[0]",
                "--",
                _arg(dest),
            ]
            result = subprocess.run(cmd, capture_output=True, check=False)
            if result.returncode == 0:
                return ConversionResult(
                    success=True,
                    output_path=dest,
                    converter_used="imagemagick",
                    warnings=warnings,
                    lossy=dest_ext in ("jpg", "jpeg"),
                )
            return ConversionResult(
                success=False,
                output_path=None,
                converter_used="imagemagick",
                error=result.stderr.decode(errors="replace")[:300],
                warnings=warnings,
            )

        return ConversionResult(
            success=False,
            output_path=None,
            converter_used=self.name,
            error="No PDF→image converter available. Install: poppler-utils (pdftoppm) or imagemagick",
        )


class ImageConverter(Converter):
    """Image ↔ image format conversion via FFmpeg (primary) or ImageMagick."""

    name = "image"
    supported = [
        ("image/", "png"),
        ("image/", "jpg"),
        ("image/", "jpeg"),
        ("image/", "webp"),
        ("image/", "gif"),
        ("image/", "bmp"),
        ("image/", "tiff"),
    ]

    _QUALITY: ClassVar[dict[str, int]] = {"low": 60, "medium": 80, "high": 95}

    def is_available(self, avail: ToolAvailability) -> bool:
        return avail.ffmpeg or avail.imagemagick

    def can_handle(self, src_mime: str, dest_ext: str) -> bool:
        # Only image→image (not PDF)
        if not src_mime.startswith("image/"):
            return False
        return dest_ext.lstrip(".").lower() in (
            "png",
            "jpg",
            "jpeg",
            "webp",
            "gif",
            "bmp",
            "tiff",
        )

    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult:
        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used="",
                error=f"Destination is an existing directory: {dest}",
            )

        avail = get_availability()
        dest_ext = dest.suffix.lstrip(".").lower()
        q = self._QUALITY.get(quality, 80)
        lossy = dest_ext in ("jpg", "jpeg", "webp")

        if avail.ffmpeg:
            cmd = ["ffmpeg", "-y", "-i", _arg(src)]
            if lossy:
                if dest_ext == "webp":
                    # WebP uses -quality 0-100 (higher = better)
                    cmd += ["-quality", str(q)]
                else:
                    # JPEG/other: ffmpeg q:v scale 1-31 (lower = better)
                    cmd += ["-q:v", str(max(1, (100 - q) // 5))]
            cmd += [_arg(dest)]
            result = subprocess.run(cmd, capture_output=True, check=False)
            if result.returncode == 0:
                return ConversionResult(
                    success=True,
                    output_path=dest,
                    converter_used="ffmpeg",
                    lossy=lossy,
                )
            logger.warning(
                "ffmpeg image convert failed: %s",
                result.stderr.decode(errors="replace")[-200:],
            )

        if avail.imagemagick:
            cmd = ["convert", _arg(src)]
            if lossy:
                cmd += ["-quality", str(q)]
            cmd += ["--", _arg(dest)]
            result = subprocess.run(cmd, capture_output=True, check=False)
            if result.returncode == 0:
                return ConversionResult(
                    success=True,
                    output_path=dest,
                    converter_used="imagemagick",
                    lossy=lossy,
                )
            return ConversionResult(
                success=False,
                output_path=None,
                converter_used="imagemagick",
                error=result.stderr.decode(errors="replace")[:300],
            )

        return ConversionResult(
            success=False,
            output_path=None,
            converter_used=self.name,
            error="No image converter available. Install: ffmpeg or imagemagick",
        )


class PDFToTextConverter(Converter):
    """PDF → plain text via pypdf or pdftotext (Poppler)."""

    name = "pdf-to-text"
    supported = [
        ("application/pdf", "txt"),
        ("application/pdf", "text"),
    ]

    def is_available(self, avail: ToolAvailability) -> bool:
        return avail.pypdf or avail.pdftotext

    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult:
        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used="",
                error=f"Destination is an existing directory: {dest}",
            )

        avail = get_availability()
        warnings: list[str] = []

        if avail.pypdf:
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(src))
                text_parts = [page.extract_text() or "" for page in reader.pages]
                extracted = "\n\n".join(text_parts)
                if extracted.strip():
                    try:
                        dest.write_text(extracted, encoding="utf-8")
                    except OSError as e:
                        return ConversionResult(
                            success=False,
                            output_path=None,
                            converter_used="pypdf",
                            error=f"Could not write {dest}: {e}",
                            warnings=warnings,
                        )
                    return ConversionResult(
                        success=True,
                        output_path=dest,
                        converter_used="pypdf",
                        metadata={"pages": len(reader.pages)},
                    )
                warnings.append(
                    "pypdf extracted no text (scanned PDF?); trying pdftotext"
                )
            except Exception as e:
                warnings.append(f"pypdf failed: {e}")

        if avail.pdftotext:
            cmd = ["pdftotext", _arg(src), _arg(dest)]
            result = subprocess.run(cmd, capture_output=True, check=False)
            if result.returncode == 0:
                return ConversionResult(
                    success=True,
                    output_path=dest,
                    converter_used="pdftotext",
                    warnings=warnings,
                )
            return ConversionResult(
                success=False,
                output_path=None,
                converter_used="pdftotext",
                error=result.stderr.decode(errors="replace")[:300],
                warnings=warnings,
            )

        return ConversionResult(
            success=False,
            output_path=None,
            converter_used=self.name,
            error="No PDF→text converter available. Install: pypdf (pip) or poppler-utils (pdftotext)",
        )


class DocumentToTextConverter(Converter):
    """DOCX/ODT/etc. → plain text or markdown via python-docx or LibreOffice."""

    name = "document-to-text"
    supported = [
        ("application/vnd.openxmlformats-officedocument.wordprocessingml", "txt"),
        ("application/vnd.openxmlformats-officedocument.wordprocessingml", "md"),
        ("application/msword", "txt"),
        ("application/msword", "md"),
    ]

    def is_available(self, avail: ToolAvailability) -> bool:
        return avail.python_docx or avail.libreoffice

    def can_handle(self, src_mime: str, dest_ext: str) -> bool:
        doc_mimes = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml",
            "application/msword",
            "application/vnd.oasis.opendocument.text",
        )
        return any(src_mime.startswith(m) for m in doc_mimes) and dest_ext.lstrip(
            "."
        ).lower() in ("txt", "md")

    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult:
        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used="",
                error=f"Destination is an existing directory: {dest}",
            )

        avail = get_availability()
        warnings: list[str] = []

        if avail.python_docx and src.suffix.lower() == ".docx":
            try:
                import docx

                doc = docx.Document(str(src))
                dest_ext = dest.suffix.lower().lstrip(".")

                if dest_ext == "md":
                    lines = [_para_to_md(para) for para in doc.paragraphs]
                else:
                    lines = [para.text for para in doc.paragraphs]

                try:
                    dest.write_text("\n".join(lines), encoding="utf-8")
                except OSError as e:
                    return ConversionResult(
                        success=False,
                        output_path=None,
                        converter_used="python-docx",
                        error=f"Could not write {dest}: {e}",
                        warnings=warnings,
                    )
                return ConversionResult(
                    success=True,
                    output_path=dest,
                    converter_used="python-docx",
                )
            except Exception as e:
                warnings.append(f"python-docx failed: {e}")

        if avail.libreoffice:
            with tempfile.TemporaryDirectory() as tmpdir:
                cmd = [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text (encoded):UTF8",
                    "--outdir",
                    tmpdir,
                    _arg(src),
                ]
                result = subprocess.run(cmd, capture_output=True, check=False)
                if result.returncode == 0:
                    out_files = list(Path(tmpdir).glob("*.txt"))
                    if out_files:
                        try:
                            out_files[0].replace(dest)
                        except OSError as e:
                            return ConversionResult(
                                success=False,
                                output_path=None,
                                converter_used="libreoffice",
                                error=f"Could not write {dest}: {e}",
                                warnings=warnings,
                            )
                        return ConversionResult(
                            success=True,
                            output_path=dest,
                            converter_used="libreoffice",
                            warnings=warnings,
                        )
                warnings.append(
                    f"libreoffice failed: {result.stderr.decode(errors='replace')[:200]}"
                )

        if (
            avail.python_docx
            and not avail.libreoffice
            and src.suffix.lower() != ".docx"
        ):
            return ConversionResult(
                success=False,
                output_path=None,
                converter_used="python-docx",
                error=f"python-docx only supports .docx files; {src.suffix!r} requires LibreOffice (not installed)",
                warnings=warnings,
            )
        return ConversionResult(
            success=False,
            output_path=None,
            converter_used=self.name,
            error="No document→text converter available. Install: python-docx (pip) or libreoffice",
            warnings=warnings,
        )


class VideoThumbnailConverter(Converter):
    """Video → JPEG thumbnail via FFmpeg."""

    name = "video-thumbnail"
    supported = [
        ("video/", "jpg"),
        ("video/", "jpeg"),
        ("video/", "png"),
    ]

    def is_available(self, avail: ToolAvailability) -> bool:
        return avail.ffmpeg

    def can_handle(self, src_mime: str, dest_ext: str) -> bool:
        return src_mime.startswith("video/") and dest_ext.lstrip(".").lower() in (
            "jpg",
            "jpeg",
            "png",
        )

    def convert(
        self,
        src: Path,
        dest: Path,
        quality: str = "medium",
        **kwargs,
    ) -> ConversionResult:
        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used="",
                error=f"Destination is an existing directory: {dest}",
            )

        avail = get_availability()
        if not avail.ffmpeg:
            return ConversionResult(
                success=False,
                output_path=None,
                converter_used=self.name,
                error="FFmpeg not available. Install: ffmpeg",
            )
        # Extract first frame (seek after -i so it works on videos of any length)
        cmd = [
            "ffmpeg",
            "-y",
            "-i",
            _arg(src),
            "-vframes",
            "1",
            _arg(dest),
        ]
        result = subprocess.run(cmd, capture_output=True, check=False)
        if result.returncode == 0:
            return ConversionResult(
                success=True,
                output_path=dest,
                converter_used="ffmpeg",
                lossy=dest.suffix.lower() in (".jpg", ".jpeg"),
            )
        return ConversionResult(
            success=False,
            output_path=None,
            converter_used="ffmpeg",
            error=result.stderr.decode(errors="replace")[-300:],
        )


# ---------------------------------------------------------------------------
# Format router
# ---------------------------------------------------------------------------

# Registry — order matters: first available converter wins
_CONVERTERS: list[Converter] = [
    PDFToImageConverter(),
    PDFToTextConverter(),
    DocumentToTextConverter(),
    VideoThumbnailConverter(),
    ImageConverter(),
]


def _detect_mime(path: Path) -> str:
    """Detect MIME type via python-magic or mimetypes fallback."""
    if get_availability().python_magic:
        try:
            import magic

            return magic.from_file(str(path), mime=True) or ""
        except Exception:
            pass
    mime, _ = mimetypes.guess_type(str(path))
    return mime or "application/octet-stream"


def find_converter(
    src: Path, dest_ext: str, avail: ToolAvailability | None = None
) -> Converter | None:
    if avail is None:
        avail = get_availability()
    src_mime = _detect_mime(src)
    for conv in _CONVERTERS:
        if conv.can_handle(src_mime, dest_ext) and conv.is_available(avail):
            return conv
    return None


def convert_file(
    src: Path,
    dest: Path,
    quality: str = "medium",
    dry_run: bool = False,
    **kwargs,
) -> ConversionResult:
    """Main entry point for file conversion.

    Args:
        src: Source file path.
        dest: Destination file path (extension determines target format).
        quality: One of 'low', 'medium', 'high'.
        dry_run: If True, return converter plan without executing.
    """
    avail = get_availability()
    src_mime = _detect_mime(src)
    dest_ext = dest.suffix.lstrip(".")

    conv = find_converter(src, dest_ext, avail)
    if conv is None:
        return ConversionResult(
            success=False,
            output_path=None,
            converter_used="none",
            error=(
                f"No converter found for {src_mime} → {dest_ext}. "
                f"Run `gptme-convert check-tools` to see available converters."
            ),
        )

    if dry_run:
        # Match the concrete converters: a directory destination cannot succeed.
        # Dry-run used to skip this check and report a possible conversion.
        if dest.is_dir():
            return ConversionResult(
                success=False,
                output_path=dest,
                converter_used=conv.name,
                error=f"Destination is an existing directory: {dest}",
            )
        return ConversionResult(
            success=True,
            output_path=dest,
            converter_used=conv.name,
            metadata={"dry_run": True, "src_mime": src_mime, "dest_ext": dest_ext},
        )

    if src.resolve() == dest.resolve():
        return ConversionResult(
            success=False,
            output_path=None,
            converter_used="none",
            error="Source and destination paths must be different",
        )

    try:
        dest.parent.mkdir(parents=True, exist_ok=True)
    except OSError as e:
        return ConversionResult(
            success=False,
            output_path=None,
            converter_used="none",
            error=f"Could not create output directory {dest.parent}: {e}",
        )
    return conv.convert(src, dest, quality=quality, **kwargs)


# ---------------------------------------------------------------------------
# Structured tool surface (auto-discovered ToolSpec)
# ---------------------------------------------------------------------------


def _execute_convert(
    code: str | None,
    args: list[str] | None,
    kwargs: dict[str, str] | None,
) -> Message:
    """Execute a file conversion from a structured tool call.

    Accepts keyword arguments ``input_path`` (required) and ``output_path``
    (required; extension selects the target format). Optional ``quality``
    (low/medium/high) and ``dry_run``.
    """
    kwargs = kwargs or {}
    src_raw = kwargs.get("input_path")
    dest_raw = kwargs.get("output_path")
    if not src_raw or not dest_raw:
        return Message(
            "system",
            "Error: both `input_path` and `output_path` are required.",
        )

    src = Path(src_raw)
    dest = Path(dest_raw)
    if not src.exists():
        return Message("system", f"Error: input file not found: {src}")

    quality = kwargs.get("quality", "medium")
    # `dry_run` may arrive as a JSON boolean or a string; normalize defensively.
    dry_run = str(kwargs.get("dry_run", "false")).strip().lower() in (
        "true",
        "1",
        "yes",
    )

    try:
        result = convert_file(src, dest, quality=quality, dry_run=dry_run)
    except OSError as exc:
        return Message("system", f"Conversion error: {exc}")
    return Message("system", result.summary())


tool = ToolSpec(
    name="convert",
    desc="Convert a file to another format using offline system tools",
    instructions=(
        "Use this tool when a workflow needs a file in a different format: "
        "render a PDF to images for visual inspection, extract text from a "
        "document for summarization, make a video thumbnail, or reformat an "
        "image. It runs fully offline, falling back gracefully when a specific "
        "converter is unavailable.\n\n"
        "Choose the destination format by the `output_path` extension (e.g. "
        "`.png`, `.jpg`, `.txt`, `.md`). Set `quality` to low/medium/high "
        "when the conversion supports it, and pass `dry_run` = true to see "
        "the converter plan first. For a new conversion, dry-run before "
        "executing to confirm the chain is available."
    ),
    execute=_execute_convert,
    parameters=[
        Parameter(
            name="input_path",
            type="string",
            description="Path to the source file to convert",
            required=True,
        ),
        Parameter(
            name="output_path",
            type="string",
            description=(
                "Path of the output file; its extension selects the target "
                "format (e.g. .png, .jpg, .txt, .md)"
            ),
            required=True,
        ),
        Parameter(
            name="quality",
            type="string",
            enum=["low", "medium", "high"],
            description="Conversion quality (default: medium)",
        ),
        Parameter(
            name="dry_run",
            type="string",
            # Keep in sync with the truthy set in `_execute_convert`.
            enum=["true", "false", "1", "yes"],
            description=(
                "If true, show the converter plan without executing. "
                "Accepts true/false/1/yes (JSON booleans also work)."
            ),
        ),
    ],
)


# ---------------------------------------------------------------------------
# CLI interface
# ---------------------------------------------------------------------------


def main(argv: list[str] | None = None) -> int:
    import argparse

    parser = argparse.ArgumentParser(
        prog="gptme-convert",
        description="Offline file format conversion for gptme agents",
    )
    sub = parser.add_subparsers(dest="cmd")

    # convert subcommand
    conv_p = sub.add_parser("convert", help="Convert a file to another format")
    conv_p.add_argument("--input", "-i", required=True, help="Input file path")
    conv_p.add_argument(
        "--output",
        "-o",
        required=True,
        help="Output file path (extension = target format)",
    )
    conv_p.add_argument(
        "--quality",
        "-q",
        choices=["low", "medium", "high"],
        default="medium",
        help="Conversion quality (default: medium)",
    )
    conv_p.add_argument(
        "--dry-run", action="store_true", help="Show converter plan without executing"
    )

    # check-tools subcommand
    sub.add_parser("check-tools", help="Show available conversion tools")

    # list-formats subcommand
    sub.add_parser("list-formats", help="List supported conversion formats")

    args = parser.parse_args(argv)

    if args.cmd == "check-tools":
        avail = get_availability()
        print("Available conversion tools:\n")
        print(avail.report())
        return 0

    if args.cmd == "list-formats":
        print("Supported conversions:")
        avail = get_availability()
        for conv in _CONVERTERS:
            status = "✓" if conv.is_available(avail) else "⚠"
            for src_mime_prefix, dest_ext in conv.supported:
                print(f"  {status} {src_mime_prefix}* → .{dest_ext}  [{conv.name}]")
        return 0

    if args.cmd == "convert":
        src = Path(args.input)
        dest = Path(args.output)
        if not src.exists():
            print(f"Error: input file not found: {src}")
            return 1
        result = convert_file(src, dest, quality=args.quality, dry_run=args.dry_run)
        print(result.summary())
        return 0 if result.success else 1

    parser.print_help()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
